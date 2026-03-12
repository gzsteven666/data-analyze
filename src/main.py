#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
业务数据分析主程序
整合所有模块，实现完整的业务洞察分析包生成流程
"""

import os
import sys
import json
import argparse
import pandas as pd
import numpy as np
from datetime import datetime
from pathlib import Path
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows

# 导入自定义模块
from data_analyzer import DataAnalyzer
from chart_generator import ChartGenerator
from infographic_generator import InfographicGenerator
from screenshot_generator import ScreenshotGenerator
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DataAnalysisPipeline:
    """数据分析流水线"""

    def __init__(self, config=None):
        """
        初始化流水线

        Args:
            config: 配置参数
        """
        self.config = config or {}
        self.analyzer = DataAnalyzer()
        self.chart_generator = None
        self.infographic_generator = None
        self.screenshot_generator = None
        self.output_root = None
        # 图表策略：auto/on/off
        self.charts_mode = self.config.get('charts_mode', 'auto')
        # 用户可指定的核心实体维度（如 医院/客户/渠道/门店/SKU 等）
        self.core_dimension = self.config.get('core_dimension')

    def resolve_output_root(self, data_path):
        """
        解析输出根目录：
        - 如果传入 output_dir，则使用传入值
        - 否则默认放在数据文件同级目录下的 outputs/
        """
        if self.config.get('output_dir'):
            return Path(self.config['output_dir']).expanduser().resolve()
        data_parent = Path(data_path).resolve().parent
        return data_parent / 'outputs'

    def setup_output_directories(self):
        """设置输出目录"""
        if self.output_root is None:
            raise ValueError("output_root 未初始化")
        directories = [
            self.output_root / 'csv',
            self.output_root / 'excel',
            self.output_root / 'figures' / 'png',
            self.output_root / 'figures' / 'svg',
            self.output_root / 'reports',
            self.output_root / 'html'
        ]

        for directory in directories:
            os.makedirs(directory, exist_ok=True)

    def init_generators(self):
        """基于输出目录初始化依赖组件"""
        if self.output_root is None:
            raise ValueError("output_root 未初始化")
        self.chart_generator = ChartGenerator(output_dir=str(self.output_root / 'figures'))
        self.infographic_generator = InfographicGenerator(
            template_dir='templates',
            output_dir=str(self.output_root / 'html')
        )
        # 截图默认关闭，若需要可后续按需初始化
        if self.config.get('enable_screenshot'):
            self.screenshot_generator = ScreenshotGenerator(
                output_dir=str(self.output_root / 'screenshots')
            )

    def should_generate_charts(self, analysis_results):
        """
        判断是否需要生成图表：
        - charts_mode == 'off' 时直接跳过
        - charts_mode == 'on' 时强制生成
        - charts_mode == 'auto' 时仅在有可视价值的板块才生成
        """
        if self.charts_mode == 'off':
            return False
        if self.charts_mode == 'on':
            return True

        # auto: 判定是否有可视价值
        numeric = analysis_results.get('数值列统计')
        cat = analysis_results.get('分类分布')
        corr = analysis_results.get('相关性矩阵')
        trend = analysis_results.get('时间趋势')
        brand = analysis_results.get('品牌份额')
        city = analysis_results.get('机会城市')
        coverage = analysis_results.get('覆盖分析')
        structure = analysis_results.get('产品结构')

        if numeric is not None and len(numeric) > 1:
            return True
        if cat is not None and not cat.empty and cat['数量'].max() > 1:
            return True
        if corr is not None:
            return True
        if trend is not None and not trend.empty and len(trend) > 2:
            return True
        if brand is not None and not brand.empty:
            return True
        if city is not None and not city.empty:
            return True
        if coverage is not None and not coverage.empty:
            return True
        if structure is not None and not structure.empty:
            return True
        return False

    def detect_data_files(self, data_dir='data'):
        """
        自动检测数据文件

        Args:
            data_dir: 数据目录

        Returns:
            list: 数据文件列表
        """
        data_files = []
        supported_extensions = ['.csv', '.xlsx', '.xls']

        if os.path.exists(data_dir):
            for file in os.listdir(data_dir):
                if any(file.endswith(ext) for ext in supported_extensions):
                    file_path = os.path.join(data_dir, file)
                    data_files.append(file_path)

        return data_files

    def build_market_narrative(
        self,
        dim_label,
        metric_label,
        concentration,
        head_names,
        tail_names,
        has_time,
        has_price,
        has_structure,
    ):
        """
        生成面向市场团队的叙事化摘要，避免学术统计口吻。
        """
        head = "、".join([str(x) for x in head_names[:3]]) if head_names else ""
        tail = "、".join([str(x) for x in tail_names[:3]]) if tail_names else ""

        if concentration:
            top3 = concentration.get("Top3占比", 0.0)
            top1 = concentration.get("Top1占比", 0.0)
            current = (
                f"当前{metric_label}主要由少数头部{dim_label}驱动，Top1约{top1:.1f}%，Top3约{top3:.1f}%。"
            )
        else:
            top3 = 0.0
            current = f"当前{metric_label}已出现明显头部与长尾分层，资源投入需要做优先级。"

        if head:
            current += f" 头部代表包括{head}。"

        opportunity = []
        if top3 >= 60:
            opportunity.append(f"头部{dim_label}是短期放量的直接抓手")
        if tail:
            opportunity.append(f"长尾中的{tail}适合作为低成本试点")
        if not opportunity:
            opportunity.append("可从覆盖不足和份额偏低的对象中筛选突破口")
        opportunity_text = "；".join(opportunity) + "。"

        strategy = (
            f"建议采用“头部深耕 + 白区试点”的双线策略：对头部{dim_label}做份额提升，对低份额高潜对象做定点突破。"
        )

        why = []
        why.append("这样做可以在不显著增加团队负担的情况下，把资源集中到最可能产生结果的对象上")
        if top3 >= 60:
            why.append("同时避免资源被低转化长尾平均摊薄")
        if not has_time:
            why.append("在缺少时间序列时，结构与覆盖策略比趋势判断更稳妥")
        why_text = "；".join(why) + "。"

        benefit = (
            f"预期收益是：更快形成可见增量、提升重点{dim_label}转化效率，并沉淀可复制打法扩展到相似对象。"
        )

        risks = []
        if not has_time:
            risks.append("缺少时间维度，短期效果评估可能滞后")
        if not has_price:
            risks.append("缺少价格维度，价格带策略判断存在盲区")
        if not has_structure:
            risks.append("缺少稳定结构字段，替代路径判断可能偏粗")
        if top3 >= 70:
            risks.append("头部集中度较高，单一对象波动会放大整体不确定性")
        if not risks:
            risks.append("当前主要风险来自执行节奏不一致")
        risk_text = "；".join(risks) + "。"

        mitigation = (
            "建议建立周节奏复盘：按对象跟踪进展、补齐关键字段（时间/价格/结构）、设置止损阈值并及时调整资源。"
        )

        return {
            "现状判断": current,
            "机会判断": opportunity_text,
            "策略建议": strategy,
            "为何现在做": why_text,
            "预期收益": benefit,
            "主要风险": risk_text,
            "风险对策": mitigation,
        }

    def humanize_dimension_label(self, column_name):
        """把原始字段名转成更自然的业务称呼。"""
        if not column_name:
            return "对象"
        label_map = {
            "品牌名称": "品牌",
            "品牌": "品牌",
            "申报企业名称": "申报企业",
            "申报企业": "申报企业",
            "企业名称": "企业",
            "厂家": "品牌",
            "生产企业": "品牌",
            "医疗机构名称": "医疗机构",
            "医院名称": "医院",
            "客户名称": "客户",
            "客户": "客户",
            "机构名称": "机构",
            "机构": "机构",
            "门店名称": "门店",
            "门店": "门店",
            "渠道名称": "渠道",
            "渠道": "渠道",
            "注册证产品名称": "产品",
            "产品名称": "产品",
            "产品": "产品",
            "目录名称": "目录",
            "产品大类": "产品大类",
            "区域": "区域",
            "地区": "区域",
        }
        if column_name in label_map:
            return label_map[column_name]
        if isinstance(column_name, str) and column_name.endswith("名称"):
            trimmed = column_name[:-2].strip()
            if trimmed:
                return trimmed
        return column_name

    def detect_missing_capabilities(self, df, price_col=None):
        """按业务概念判断数据缺口，避免把字段别名误报成缺失。"""
        concept_map = {
            "时间维度": ["时间", "日期", "月份", "月", "周", "季度", "年月"],
            "渠道维度": ["渠道", "渠道名称", "销售渠道", "渠道类型"],
            "价格字段": [price_col] if price_col else ["价格", "单价", "中选价格"],
        }
        missing = []
        for label, candidates in concept_map.items():
            valid_candidates = [cand for cand in candidates if cand]
            if not any(cand in df.columns for cand in valid_candidates):
                missing.append(label)
        return missing

    def build_strategic_insights(self, analysis_results, blueprint):
        """把报告蓝图转成统一的 JSON 洞察结构，供导出与前端共用。"""
        insights = {
            "report_style": {
                "reference_expert": "Richard Rumelt",
                "why": "采用“诊断-指导方针-协调行动”的结构，避免把洞察写成泛泛罗列的机会清单。"
            },
            "executive_summary": blueprint.get("executive_summary", []),
            "current_state": blueprint.get("snapshot", []),
            "core_diagnosis": {
                "primary_issue": blueprint.get("primary_issue", ""),
                "supporting_points": blueprint.get("diagnosis_points", []),
            },
            "model_analysis": [
                {
                    "title": section.get("title", ""),
                    "judgment": section.get("judgment", ""),
                    "evidence": section.get("evidence", ""),
                    "implication": section.get("implication", ""),
                    "action": section.get("action", ""),
                }
                for section in blueprint.get("model_sections", [])
            ],
            "guiding_policies": blueprint.get("guiding_policies", []),
            "priority_focus": blueprint.get("priority_focus", {}),
            "action_plan_90d": blueprint.get("action_plan", []),
            "do_not_do": blueprint.get("do_not_do", []),
            "risk_controls": blueprint.get("risk_controls", []),
        }

        target_profile = blueprint.get("target_profile")
        if target_profile:
            insights["市场地位"] = {
                "目标对象": target_profile.get("name"),
                "排名": target_profile.get("rank"),
                "份额": f"{target_profile.get('share', 0.0):.2f}%",
                "领先对象": target_profile.get("leader_name"),
                "领先差距": f"{target_profile.get('leader_gap', 0.0):.2f}pct",
            }

        channel_focus = blueprint.get("channel_focus")
        if channel_focus:
            insights["机会渠道"] = {
                "最佳机会": channel_focus.get("name"),
                "容量": f"{channel_focus.get('total', 0.0):,.0f}",
                "当前份额": f"{channel_focus.get('share', 0.0):.2f}%",
            }

        city_focus = blueprint.get("city_focus")
        if city_focus:
            insights["机会城市"] = {
                "最佳机会": city_focus.get("name"),
                "容量": f"{city_focus.get('total', 0.0):,.0f}",
                "当前份额": f"{city_focus.get('share', 0.0):.2f}%",
            }

        focus_cards = []
        for item in blueprint.get("diagnosis_points", [])[:3]:
            focus_cards.append({
                "icon": "fas fa-magnifying-glass-chart",
                "title": "核心诊断",
                "content": item,
            })
        if focus_cards:
            insights["通用洞察"] = focus_cards

        opportunity_cards = []
        for action in blueprint.get("action_plan", [])[:3]:
            opportunity_cards.append({
                "icon": "fas fa-bullseye",
                "title": "行动建议",
                "description": action,
            })
        if opportunity_cards:
            insights["机会建议"] = opportunity_cards

        return insights

    def analyze_data(self, data_path, sheet_name=None, company_name=None):
        """
        执行数据分析

        Args:
            data_path: 数据文件路径
            sheet_name: Excel工作表名称
            company_name: 目标公司名称

        Returns:
            dict: 分析结果
        """
        print(f"开始分析数据: {data_path}")

        # 1. 加载数据
        df = self.analyzer.load_data(data_path, sheet_name)
        df = self.analyzer.preprocess_data(df)
        print(f"数据加载完成，共{len(df)}条记录，{len(df.columns)}个字段")

        # 2. 数据体检
        health_report = self.analyzer.data_health_check(df)
        print("数据体检完成")

        # 3. 过滤留置针数据（如果适用）
        if '品种名称' in df.columns:
            df_filtered = self.analyzer.filter_iv_catheter(df)
            print(f"过滤后留置针数据: {len(df_filtered)}条记录")
        else:
            df_filtered = df

        # 4. 执行综合分析
        analysis_results = self.analyzer.create_comprehensive_analysis(
            df_filtered, company_name, self.config
        )
        print("综合分析完成")

        # 5. 生成统一的战略蓝图与洞察，保证 Word/JSON/HTML 口径一致
        blueprint = self.build_strategic_report_blueprint(df_filtered, analysis_results)
        insights = self.build_strategic_insights(analysis_results, blueprint)

        return {
            'raw_data': df,
            'filtered_data': df_filtered,
            'health_report': health_report,
            'analysis_results': analysis_results,
            'blueprint': blueprint,
            'insights': insights
        }

    def generate_charts(self, analysis_results, company_name=None):
        """
        生成可视化图表

        Args:
            analysis_results: 分析结果

        Returns:
            dict: 图表路径
        """
        print("开始生成图表...")
        chart_paths = self.chart_generator.generate_all_charts(
            analysis_results, company_name=company_name or ''
        )
        print(f"图表生成完成，共{len(chart_paths)}个图表")
        return chart_paths

    def export_data(self, analysis_results, insights):
        """
        导出分析数据

        Args:
            analysis_results: 分析结果
            insights: 洞察分析
        """
        print("开始导出数据...")

        def make_json_safe(value):
            if isinstance(value, dict):
                return {str(k): make_json_safe(v) for k, v in value.items()}
            if isinstance(value, list):
                return [make_json_safe(v) for v in value]
            if isinstance(value, tuple):
                return [make_json_safe(v) for v in value]
            if isinstance(value, (np.integer, np.floating)):
                return value.item()
            if isinstance(value, pd.Timestamp):
                return value.isoformat()
            return value

        csv_dir = self.output_root / 'csv'
        # 导出CSV文件
        csv_exports = {
            '字段概览': analysis_results.get('字段概览'),
            '数值列统计': analysis_results.get('数值列统计'),
            '分类分布': analysis_results.get('分类分布'),
            '相关性矩阵': analysis_results.get('相关性矩阵'),
            '时间趋势': analysis_results.get('时间趋势'),
            '核心维度分布': analysis_results.get('核心维度分布'),
            '城市分布': analysis_results.get('城市分布'),
            '渠道分布': analysis_results.get('渠道分布'),
            '城市品牌分布': analysis_results.get('城市品牌分布'),
            '城市Top3': analysis_results.get('城市Top3'),
            '城市白区': analysis_results.get('城市白区'),
            '机会优先级_城市': analysis_results.get('机会优先级_城市'),
            '机会城市': analysis_results.get('机会城市'),
            '渠道白区': analysis_results.get('渠道白区'),
            '机会优先级_渠道': analysis_results.get('机会优先级_渠道'),
            '重点实体白区': analysis_results.get('重点实体白区'),
            '机会优先级_重点实体': analysis_results.get('机会优先级_重点实体'),
            '目录分布': analysis_results.get('目录分布'),
            '大类分布': analysis_results.get('大类分布'),
            '产品结构': analysis_results.get('产品结构'),
            '覆盖分析': analysis_results.get('覆盖分析'),
            '重点实体TOP': analysis_results.get('重点实体TOP'),
            '产品分布': analysis_results.get('产品分布'),
            '品牌产品分布': analysis_results.get('品牌产品分布'),
            '城市产品分布': analysis_results.get('城市产品分布'),
            '品牌产品Top': analysis_results.get('品牌产品Top'),
            # 兼容旧版输出键，避免依赖旧文件名的链路立刻失效
            '机会医院': analysis_results.get('机会医院'),
            '医院TOP': analysis_results.get('医院TOP'),
            '医院白院': analysis_results.get('医院白院'),
            '机会优先级_医院': analysis_results.get('机会优先级_医院')
        }

        for name, data in csv_exports.items():
            if data is not None and not data.empty:
                filepath = csv_dir / f"{name}.csv"
                data.to_csv(filepath, index=False, encoding='utf-8-sig')
                print(f"导出CSV: {filepath}")

        # 导出Excel汇总表（带目录，移除字段概览/数值统计）
        excel_path = self.output_root / "excel/数据分析汇总.xlsx"
        wb = Workbook()
        toc = wb.active
        toc.title = "目录"
        toc.append(["名称", "跳转"])

        def add_toc_row(name):
            row = toc.max_row + 1
            toc.cell(row=row, column=1, value=name)
            cell = toc.cell(row=row, column=2, value="点击跳转")
            cell.hyperlink = f"#{name}!A1"
            cell.style = "Hyperlink"

        exclude_excel = {"字段概览", "数值列统计"}
        for name, data in csv_exports.items():
            if name in exclude_excel or data is None or data.empty:
                continue
            ws = wb.create_sheet(name)
            for r in dataframe_to_rows(data, index=False, header=True):
                ws.append(r)
            add_toc_row(name)
            back = ws.cell(row=ws.max_row + 2, column=1)
            back.value = "返回目录"
            back.hyperlink = "#目录!A1"
            back.style = "Hyperlink"

        if "Sheet" in wb.sheetnames:
            wb.remove(wb["Sheet"])
        wb.save(excel_path)
        print(f"导出Excel: {excel_path}")

        # 导出洞察分析
        insights_path = self.output_root / "reports/洞察分析.json"
        with open(insights_path, 'w', encoding='utf-8') as f:
            json.dump(make_json_safe(insights), f, ensure_ascii=False, indent=2)
        print(f"导出洞察: {insights_path}")

    def generate_infographic(self, analysis_results, chart_paths, insights, company_name=None):
        """
        生成信息图HTML

        Args:
            analysis_results: 分析结果
            chart_paths: 图表路径
            insights: 洞察分析
            company_name: 公司名称

        Returns:
            str: HTML文件路径
        """
        print("生成信息图HTML...")
        html_content = self.infographic_generator.generate_infographic_html(
            analysis_results, chart_paths, insights, company_name
        )

        html_path = self.infographic_generator.save_infographic(html_content)
        print(f"信息图HTML已生成: {html_path}")
        return html_path

    def generate_screenshot(self, html_path, output_name=None):
        """按需生成截图（默认关闭）"""
        if not self.config.get('enable_screenshot') or not self.screenshot_generator:
            return None
        print("生成截图...")
        screenshot_path = self.screenshot_generator.generate_screenshot_sync(
            html_path, output_name
        )
        print(f"截图已生成: {screenshot_path}")
        return screenshot_path

    def create_chart_gallery(self, chart_paths):
        """生成图表汇总 HTML，方便一次性查看/下载"""
        if not chart_paths:
            return None
        html_dir = self.output_root / 'html'
        html_dir.mkdir(parents=True, exist_ok=True)
        mapping = {
            'core_share': '核心维度份额',
            'city_share': '城市分布',
            'category_share': '目录份额',
            'major_share': '产品大类分布',
            'coverage': '覆盖与单实体均量',
            'city_opportunities': '机会城市',
            'product_structure': '产品结构',
            'time_trend': '时间趋势',
            'correlation_heatmap': '相关性',
            'categorical_topn': '分类TopN',
            'numeric_overview': '数值概览',
        }
        items = []
        from pathlib import Path
        for key, val in chart_paths.items():
            png = None
            if isinstance(val, dict):
                png = val.get('png')
            elif isinstance(val, (list, tuple)) and len(val) > 0:
                png = val[0]
            if png and Path(png).exists():
                items.append((mapping.get(key, key), png))
        if not items:
            return None

        html_parts = [
            "<!DOCTYPE html>",
            "<html lang='zh-CN'>",
            "<head>",
            "<meta charset='UTF-8' />",
            "<title>图表汇总</title>",
            "<style>",
            "body { font-family: Arial, sans-serif; margin: 24px; background: #f7f7f7; }",
            "h1 { margin-bottom: 12px; }",
            ".card { background: #fff; padding: 16px; margin-bottom: 16px; box-shadow: 0 2px 6px rgba(0,0,0,0.08); border-radius: 8px; }",
            ".card h2 { margin: 0 0 8px 0; font-size: 18px; }",
            ".card img { max-width: 100%; height: auto; border: 1px solid #e5e5e5; border-radius: 4px; }",
            "</style>",
            "</head>",
            "<body>",
            "<h1>图表汇总</h1>",
        ]
        for title, png in items:
            html_parts.append("<div class='card'>")
            html_parts.append(f"<h2>{title}</h2>")
            html_parts.append(f"<img src='file://{png}' alt='{title}' />")
            html_parts.append("</div>")
        html_parts.append("</body></html>")
        out_path = html_dir / "charts_gallery.html"
        out_path.write_text("\n".join(html_parts), encoding="utf-8")
        print(f"图表汇总已生成: {out_path}")
        return str(out_path)

    def format_entities(self, names, max_n=5):
        """格式化实体名称列表。"""
        cleaned = [
            str(n).strip()
            for n in names
            if pd.notna(n) and str(n).strip() and str(n).strip().lower() not in {'nan', 'none', 'null', 'na', 'n/a'}
        ]
        if not cleaned:
            return ""
        subset = cleaned[:max_n]
        return "、".join(subset) + (" 等" if len(cleaned) > max_n else "")

    def _pick_entity_priority(self, analysis_results):
        entity_priority = analysis_results.get('机会优先级_重点实体')
        if entity_priority is None or entity_priority.empty:
            entity_priority = analysis_results.get('机会优先级_医院')
        return entity_priority

    def _detect_entity_name_col(self, df):
        if df is None:
            return None
        for cand in ['医院名称', '医疗机构名称', '客户名称', '客户', '机构名称', '机构', '门店名称', '门店', '重点实体']:
            if cand in df.columns:
                return cand
        return None

    def _detect_total_col(self, df):
        if df is None:
            return None
        for cand in ['重点实体总量', '医院总量', '客户总量', '机构总量', '门店总量', '城市总量', '总容量', '总量']:
            if cand in df.columns:
                return cand
        for col in df.columns:
            if '总量' in str(col) or '容量' in str(col):
                return col
        return None

    def build_strategic_report_blueprint(self, filtered_df, analysis_results):
        """
        采用“诊断 -> 指导方针 -> 协调行动”框架组织洞察报告。
        目标不是堆砌分析结果，而是把真正需要管理层决策的问题说清楚。
        """
        core_metric = analysis_results.get('核心指标列') or '核心指标'
        core_dim = analysis_results.get('核心维度') or '核心维度'
        dim_label = self.humanize_dimension_label(core_dim)
        concentration = analysis_results.get('集中度分析') or {}
        top1 = float(concentration.get('Top1占比', 0.0) or 0.0)
        top3 = float(concentration.get('Top3占比', 0.0) or 0.0)
        top5 = float(concentration.get('Top5占比', 0.0) or 0.0)
        target_object = analysis_results.get('目标对象')

        head_names = analysis_results.get('头部名单') or []
        tail_names = analysis_results.get('尾部名单') or []
        head_names_fmt = self.format_entities(head_names)
        tail_names_fmt = self.format_entities(tail_names)

        core_share = analysis_results.get('核心维度分布')
        city_priority = analysis_results.get('机会优先级_城市')
        channel_priority = analysis_results.get('机会优先级_渠道')
        channel_share = analysis_results.get('渠道分布')
        product_dist = analysis_results.get('产品分布')
        entity_priority = self._pick_entity_priority(analysis_results)
        entity_name_col = self._detect_entity_name_col(entity_priority)
        entity_total_col = self._detect_total_col(entity_priority)

        price_info = analysis_results.get('价格分析') or {}
        price_col = price_info.get('价格列')
        price_insight = None
        if price_info:
            corr = price_info.get('相关系数')
            relation = price_info.get('相关描述')
            if corr is not None and relation and price_col:
                price_insight = f"{price_col} 与 {core_metric}的关系判断为“{relation}”（相关系数约 {corr:.2f}）"

        structure_summary = analysis_results.get('结构概览') or {}
        structure_insight = None
        if structure_summary:
            main_type = structure_summary.get('主类型')
            main_share = structure_summary.get('主类型占比')
            others = structure_summary.get('其他占比')
            if main_type is not None and main_share is not None:
                structure_insight = f"当前结构主要由 {main_type} 承担（约 {main_share:.1f}%）"
                if others:
                    structure_insight += f"，其余结构为 {others}"

        trend_summary = analysis_results.get('时间趋势摘要') or {}
        time_trend_summary = None
        trend_direction = None
        if trend_summary:
            start_v = trend_summary.get('起始')
            end_v = trend_summary.get('当前')
            trend_direction = trend_summary.get('方向')
            if start_v is not None and end_v is not None and trend_direction:
                time_trend_summary = f"整体时间趋势呈{trend_direction}态势（{start_v:,.1f} -> {end_v:,.1f}）"

        numeric_stats = analysis_results.get('数值列统计')
        core_metric_range = None
        if numeric_stats is not None and not numeric_stats.empty and core_metric in numeric_stats['字段'].values:
            row = numeric_stats[numeric_stats['字段'] == core_metric].iloc[0]
            core_metric_range = f"{core_metric}的区间大致在 {row['最小值']:,.2f} 到 {row['最大值']:,.2f}"

        missing_fields = self.detect_missing_capabilities(filtered_df, price_col=price_col)

        target_profile = None
        if (
            target_object
            and core_share is not None
            and not core_share.empty
            and core_dim in core_share.columns
            and core_metric in core_share.columns
            and target_object in core_share[core_dim].values
        ):
            rank_source = core_share.sort_values(core_metric, ascending=False).reset_index(drop=True)
            target_row = rank_source[rank_source[core_dim] == target_object].iloc[0]
            target_rank = int(rank_source[rank_source[core_dim] == target_object].index[0] + 1)
            leader_row = rank_source.iloc[0]
            target_profile = {
                'name': target_object,
                'rank': target_rank,
                'share': float(target_row.get('份额(%)', 0.0) or 0.0),
                'volume': float(target_row.get(core_metric, 0.0) or 0.0),
                'leader_name': leader_row.get(core_dim, ''),
                'leader_share': float(leader_row.get('份额(%)', 0.0) or 0.0),
                'leader_gap': float((leader_row.get('份额(%)', 0.0) or 0.0) - (target_row.get('份额(%)', 0.0) or 0.0)),
            }

        city_focus = None
        if city_priority is not None and not city_priority.empty and '城市' in city_priority.columns:
            top_city = city_priority.iloc[0]
            city_focus = {
                'name': top_city['城市'],
                'share': float(top_city.get('目标品牌份额(%)', 0.0) or 0.0),
                'total': float(top_city.get('城市总量', 0.0) or 0.0),
                'score': float(top_city.get('综合优先级分', 0.0) or 0.0),
                'reason': top_city.get('优先级理由', '')
            }

        entity_focus = None
        if entity_priority is not None and not entity_priority.empty and entity_name_col:
            top_entity = entity_priority.iloc[0]
            entity_focus = {
                'name': top_entity.get(entity_name_col, '重点实体'),
                'share': float(top_entity.get('目标品牌份额(%)', 0.0) or 0.0),
                'total': float(top_entity.get(entity_total_col, 0.0) or 0.0) if entity_total_col else 0.0,
                'score': float(top_entity.get('综合优先级分', 0.0) or 0.0),
                'reason': top_entity.get('优先级理由', '')
            }

        channel_focus = None
        if channel_priority is not None and not channel_priority.empty:
            channel_name_col = next(
                (cand for cand in ['渠道', '渠道名称', '销售渠道', '渠道类型'] if cand in channel_priority.columns),
                None
            )
            if channel_name_col:
                top_channel = channel_priority.iloc[0]
                channel_focus = {
                    'name': top_channel.get(channel_name_col, '渠道'),
                    'share': float(top_channel.get('目标品牌份额(%)', 0.0) or 0.0),
                    'total': float(top_channel.get('渠道总量', 0.0) or 0.0),
                    'score': float(top_channel.get('综合优先级分', 0.0) or 0.0),
                    'reason': top_channel.get('优先级理由', ''),
                }
        elif channel_share is not None and not channel_share.empty:
            top_channel = channel_share.iloc[0]
            channel_focus = {
                'name': top_channel.get('渠道', '渠道'),
                'share': float(top_channel.get('份额(%)', 0.0) or 0.0),
                'total': float(top_channel.get(core_metric, 0.0) or 0.0),
                'score': 0.0,
                'reason': '',
            }

        product_focus = None
        if product_dist is not None and not product_dist.empty and core_metric in product_dist.columns:
            product_rank = product_dist.sort_values(core_metric, ascending=False).reset_index(drop=True)
            total_product = float(product_rank[core_metric].sum() or 0.0)
            top_product = product_rank.iloc[0]
            tail_product = product_rank.iloc[-1]
            product_focus = {
                'leader_name': top_product.get('注册证产品名称', '主力产品'),
                'leader_share': float(top_product.get(core_metric, 0.0) / total_product * 100) if total_product > 0 else 0.0,
                'tail_name': tail_product.get('注册证产品名称', '尾部产品'),
                'tail_share': float(tail_product.get(core_metric, 0.0) / total_product * 100) if total_product > 0 else 0.0,
            }

        if target_profile and city_focus and target_profile['leader_gap'] >= 10:
            primary_issue = (
                f"当前主要矛盾不是市场没有需求，而是{target_profile['name']}在整体{dim_label}格局中仍落后头部，"
                f"同时高容量低份额区域尚未形成明确进攻顺序；如果继续平均铺资源，差距很难缩小。"
            )
        elif top3 >= 65 and city_focus:
            primary_issue = (
                f"当前主要矛盾不是缺少机会，而是{core_metric}已高度集中在少数头部{dim_label}，"
                f"同时高容量低份额区域仍未被优先级管理；如果资源继续平均铺开，增量会被稀释。"
            )
        elif top3 >= 65:
            primary_issue = (
                f"当前主要矛盾不是市场盘子不够，而是增长过度依赖少数头部{dim_label}；"
                f"只要资源分散，短期结果就会明显变差。"
            )
        elif city_focus:
            primary_issue = (
                f"当前主要矛盾不是需求不足，而是高容量低份额区域尚未形成明确的进攻顺序；"
                f"先后顺序不清会导致投入分散、组织感觉忙但结果不集中。"
            )
        elif trend_direction == '下降':
            primary_issue = (
                "当前主要矛盾不是继续扩更多点，而是先修复下滑环节，"
                "把有限资源集中到最容易形成确定性结果的对象上。"
            )
        else:
            primary_issue = (
                "当前主要矛盾不是分析维度不够，而是要先把最能转化的对象和动作顺序排出来，"
                "避免在缺乏取舍的情况下同时做太多事。"
            )

        executive_summary = [
            f"这份数据首先说明，当前业务并非没有机会，而是机会和结果都呈现明显分层：核心观察维度是“{dim_label}”，核心指标是“{core_metric}”。",
            primary_issue,
            "因此，本报告不建议把资源平均分配，而是建议采用“头部稳盘 + 白区试点 + 节点复盘”的组合打法。"
        ]

        snapshot = [
            f"本次分析覆盖 {len(filtered_df)} 条记录、{len(filtered_df.columns)} 个字段，当前主要观察口径为“{dim_label}”和“{core_metric}”。",
        ]
        if target_profile:
            snapshot.append(
                f"以 {target_profile['name']} 为目标对象看，当前排名第 {target_profile['rank']}，份额约 {target_profile['share']:.1f}%，"
                f"落后第一名 {target_profile['leader_name']} 约 {target_profile['leader_gap']:.1f} 个百分点。"
            )
        if core_metric_range:
            snapshot.append(core_metric_range + "，说明对象体量差异较大，应按分层而不是统一动作管理。")
        if channel_focus and channel_share is not None and not channel_share.empty:
            if channel_priority is not None and not channel_priority.empty:
                snapshot.append(
                    f"渠道上，{channel_focus['name']} 这类通道容量约 {channel_focus['total']:,.0f}，当前目标份额约 {channel_focus['share']:.1f}%，值得单独管理而不是与其他渠道混在一起推进。"
                )
            else:
                snapshot.append(
                    f"渠道上，{channel_focus['name']} 已承载约 {channel_focus['total']:,.0f} 的体量，是当前最重要的结果通道之一。"
                )
        if product_focus:
            snapshot.append(
                f"产品上，{product_focus['leader_name']} 是当前主力（约 {product_focus['leader_share']:.1f}%），"
                f"{product_focus['tail_name']} 贡献较低（约 {product_focus['tail_share']:.1f}%），更适合作为结构补充而非统一主推。"
            )
        if time_trend_summary:
            snapshot.append("从时间维度看，" + time_trend_summary + "。")
        if structure_insight:
            snapshot.append("从结构维度看，" + structure_insight + "。")

        diagnosis_points = []
        if concentration:
            if top3 >= 65:
                diagnosis_points.append(
                    f"{core_metric}高度集中：Top1 约 {top1:.1f}%，Top3 约 {top3:.1f}%，Top5 约 {top5:.1f}%。这意味着真正的任务不是广撒网，而是管好少数关键{dim_label}。"
                )
            else:
                diagnosis_points.append(
                    f"{core_metric}分布相对分散：Top3 约 {top3:.1f}%，说明增长更依赖优先级管理和标准打法复制，而不是单一头部拉动。"
                )
        if target_profile:
            diagnosis_points.append(
                f"{target_profile['name']} 当前处于第 {target_profile['rank']} 位，份额约 {target_profile['share']:.1f}%。如果目标是继续追份额，首要问题不是再铺更多点，而是先在最容易拉开差距的对象上追平领先者。"
            )
        if city_focus:
            diagnosis_points.append(
                f"区域上存在明确白区：如 {city_focus['name']}，容量约 {city_focus['total']:,.0f}，当前目标份额仅 {city_focus['share']:.1f}%，说明问题更像渗透不足而不是市场不足。"
            )
        if channel_focus and channel_priority is not None and not channel_priority.empty:
            diagnosis_points.append(
                f"渠道上存在投入顺序问题：{channel_focus['name']} 渠道容量约 {channel_focus['total']:,.0f}，当前目标份额仅 {channel_focus['share']:.1f}%，说明资源配置与高潜渠道并未完全对齐。"
            )
        if entity_focus:
            diagnosis_points.append(
                f"在重点实体层，{entity_focus['name']} 的综合优先级较高（{entity_focus['score']:.1f} 分），适合作为近期验证打法的样板点。"
            )
        if product_focus:
            diagnosis_points.append(
                f"产品结构存在主次之分：{product_focus['leader_name']} 承担主要结果，而 {product_focus['tail_name']} 占比仅 {product_focus['tail_share']:.1f}%，不宜在所有渠道同步重推。"
            )
        if price_insight:
            diagnosis_points.append("价格层面，" + price_insight + "，因此价格策略应按对象分层，而不宜统一处理。")
        if missing_fields:
            diagnosis_points.append(
                "当前数据仍缺少 {} 等关键信息，因此报告中的判断更偏向结构诊断和优先级排序，而不是完整因果归因。".format(
                    "、".join(missing_fields)
                )
            )

        model_sections = []
        if concentration:
            model_sections.append({
                'title': '模型一：集中度与资源聚焦',
                'judgment': (
                    f"当前增长主要由少数头部{dim_label}决定，管理重点应从“多点铺开”切换为“有限目标的深耕”。"
                    if top3 >= 65 else
                    "当前市场并未被绝对头部锁死，更关键的是建立可复制的分层打法。"
                ),
                'evidence': f"Top1 {top1:.1f}% / Top3 {top3:.1f}% / Top5 {top5:.1f}%；头部对象包括 {head_names_fmt or '暂无明显头部对象'}。",
                'implication': f"头部{dim_label}应承担短期结果，尾部只保留试点用途，不建议在同一阶段同时对所有对象做重投入。",
                'action': f"先锁定最关键的头部{dim_label}作为守盘对象，再把剩余资源给到高优先级白区。"
            })
        if target_profile or city_focus or entity_focus:
            evidence_parts = []
            if target_profile:
                evidence_parts.append(
                    f"{target_profile['name']} 当前排名第 {target_profile['rank']}，份额 {target_profile['share']:.1f}%，落后第一名 {target_profile['leader_name']} 约 {target_profile['leader_gap']:.1f}pct"
                )
            if city_focus:
                evidence_parts.append(
                    f"城市优先项 {city_focus['name']}：容量 {city_focus['total']:,.0f}、当前份额 {city_focus['share']:.1f}%、综合分 {city_focus['score']:.1f}"
                )
            if entity_focus:
                evidence_parts.append(
                    f"重点实体优先项 {entity_focus['name']}：容量 {entity_focus['total']:,.0f}、当前份额 {entity_focus['share']:.1f}%、综合分 {entity_focus['score']:.1f}"
                )
            model_sections.append({
                'title': '模型二：份额缺口 × 容量',
                'judgment': "高容量低份额对象值得优先进入，但前提是按优先级推进，而不是同时全面进攻。",
                'evidence': "；".join(evidence_parts),
                'implication': "白区不等于全面铺货。正确做法是先拿下最容易产生示范效应的城市或重点实体，再决定是否扩面。",
                'action': "优先打前三个高容量低份额对象，要求每个对象都有负责人、节奏和止损条件。"
            })
        if channel_focus or product_focus or structure_insight or time_trend_summary or price_insight:
            evidence_parts = []
            if channel_focus:
                if channel_priority is not None and not channel_priority.empty:
                    evidence_parts.append(
                        f"渠道优先项 {channel_focus['name']}：容量 {channel_focus['total']:,.0f}、当前份额 {channel_focus['share']:.1f}%、综合分 {channel_focus['score']:.1f}"
                    )
                else:
                    evidence_parts.append(
                        f"主力渠道 {channel_focus['name']} 当前承载体量约 {channel_focus['total']:,.0f}，份额约 {channel_focus['share']:.1f}%"
                    )
            if product_focus:
                evidence_parts.append(
                    f"主力产品 {product_focus['leader_name']} 约占 {product_focus['leader_share']:.1f}%，尾部产品 {product_focus['tail_name']} 约占 {product_focus['tail_share']:.1f}%"
                )
            evidence_parts.extend([x for x in [structure_insight, time_trend_summary, price_insight] if x])
            model_sections.append({
                'title': '模型三：渠道与产品的取舍',
                'judgment': "决定动作强度的不是单一结果数字，而是渠道是否承接得住、产品是否值得主推，以及当前结构是否支持扩张。",
                'evidence': "；".join(evidence_parts),
                'implication': "即使看见机会，也应先判断当前渠道和产品组合是否承受得住扩张，否则容易出现短期冲量、后续失速。",
                'action': "把主力产品放在最有把握的渠道里继续放大，尾部产品只在局部场景做验证，不做全面铺货。"
            })
        if missing_fields or tail_names_fmt:
            reality_parts = []
            if tail_names_fmt:
                reality_parts.append(f"长尾对象如 {tail_names_fmt} 更适合低成本试点")
            if missing_fields:
                reality_parts.append(f"缺少字段：{'、'.join(missing_fields)}")
            model_sections.append({
                'title': '模型四：现实约束校验',
                'judgment': "报告建议必须建立在现实约束上，凡是当前数据无法验证的判断，都不应直接升级成大规模投入计划。",
                'evidence': "；".join(reality_parts) if reality_parts else "当前数据对因果归因支持有限。",
                'implication': "先做可验证、可止损的小动作，再决定是否扩张，是比一次性重投入更稳妥的管理方式。",
                'action': "把无法验证的判断降级为试点假设，用 4-6 周结果决定是否继续投入。"
            })
        for idx, section in enumerate(model_sections, start=1):
            if '：' in section['title']:
                section['title'] = f"模型{idx}：" + section['title'].split('：', 1)[1]

        guiding_policies = []
        if top3 >= 60:
            guiding_policies.append(f"优先守住头部{dim_label}，把短期结果建立在最确定的对象上，而不是平均分配资源。")
        else:
            guiding_policies.append(f"按分层而不是按直觉管理{dim_label}，先打磨一套可复制打法，再做扩面。")
        if target_profile:
            guiding_policies.append(
                f"围绕 {target_profile['name']} 追份额时，先追最容易缩小差距的城市、渠道和重点实体，不以“全面覆盖”作为阶段目标。"
            )
        if city_focus:
            guiding_policies.append("区域动作按“容量 × 份额缺口”排序推进，不按行政区平均覆盖。")
        if channel_focus and channel_priority is not None and not channel_priority.empty:
            guiding_policies.append("渠道资源按“高容量低份额优先”排序，而不是按历史惯性平均铺预算。")
        if entity_focus:
            guiding_policies.append("重点实体推进采用样板点策略，先验证进入和转化，再决定是否复制。")
        if product_focus:
            guiding_policies.append(
                f"产品策略上，{product_focus['leader_name']} 继续承担主攻，{product_focus['tail_name']} 只在适配场景做组合测试。"
            )
        if price_insight:
            guiding_policies.append("价格策略按对象分层，不做一次性统一调价。")
        guiding_policies.append("凡当前数据不能证明的判断，只允许小范围试点，不允许大面积押注。")

        priority_focus = {
            'city': [],
            'channel': [],
            'entity': []
        }
        if city_priority is not None and not city_priority.empty and '城市' in city_priority.columns:
            for _, row in city_priority.head(3).iterrows():
                priority_focus['city'].append(
                    f"{row['城市']}：综合分 {float(row.get('综合优先级分', 0.0)):.1f}，当前份额 {float(row.get('目标品牌份额(%)', 0.0)):.1f}%，容量 {float(row.get('城市总量', 0.0)):.0f}。"
                )
        if channel_priority is not None and not channel_priority.empty:
            name_col = next((cand for cand in ['渠道', '渠道名称', '销售渠道', '渠道类型'] if cand in channel_priority.columns), None)
            if name_col:
                for _, row in channel_priority.head(3).iterrows():
                    priority_focus['channel'].append(
                        f"{row[name_col]}：综合分 {float(row.get('综合优先级分', 0.0)):.1f}，当前份额 {float(row.get('目标品牌份额(%)', 0.0)):.1f}%，容量 {float(row.get('渠道总量', 0.0)):.0f}。"
                    )
        if entity_priority is not None and not entity_priority.empty and entity_name_col:
            for _, row in entity_priority.head(3).iterrows():
                priority_focus['entity'].append(
                    f"{row.get(entity_name_col, '重点实体')}：综合分 {float(row.get('综合优先级分', 0.0)):.1f}，当前份额 {float(row.get('目标品牌份额(%)', 0.0)):.1f}%，容量 {float(row.get(entity_total_col, 0.0)):.0f}。"
                )

        action_plan = [
            "近30天：只锁定一小批头部对象和高优先机会对象，逐个明确负责人、目标口径和每周复盘节奏，不在同一阶段同时追求新增、提价、上新三件事。"
        ]
        if city_focus:
            action_plan.append(
                f"30-60天：把 {city_focus['name']} 这类高容量低份额城市作为区域试点，优先验证进入效率和资源消耗，而不是一开始就全面铺开。"
            )
        if channel_focus and channel_priority is not None and not channel_priority.empty:
            action_plan.append(
                f"30-60天：把 {channel_focus['name']} 这类高容量低份额渠道作为单独战役管理，明确预算边界、主推产品和复盘指标。"
            )
        if entity_focus:
            action_plan.append(
                f"30-60天：以 {entity_focus['name']} 这类高优先重点实体做样板点，要求形成可复盘案例、阻塞点清单和复制条件。"
            )
        if product_focus:
            action_plan.append(
                f"60-90天：围绕 {product_focus['leader_name']} 继续放大确定性结果，{product_focus['tail_name']} 仅在适配区域或渠道做低预算验证。"
            )
        elif structure_insight:
            action_plan.append("60-90天：先放大当前主结构的确定性收益，再用小预算测试次结构，防止一开始就把资源分散到过多结构上。")
        action_plan.append(
            "设置止损条件：若4-6周内重点对象的份额、覆盖或转化没有出现明确改善，就暂停扩面，转向下一优先级对象。"
        )
        if missing_fields:
            action_plan.append(
                "同步补数：在后续采集里补齐 {}，否则下一轮只能继续做结构判断，无法更深入验证策略效果。".format(
                    "、".join(missing_fields)
                )
            )

        do_not_do = [
            "不要把所有城市、渠道和重点实体同时列为优先级对象，否则组织动作会先被摊薄。",
            "不要把白区自动当成必投区域，高容量低份额只能说明值得验证，不能替代一线可进入性判断。",
        ]
        if product_focus:
            do_not_do.append(
                f"不要把 {product_focus['tail_name']} 这类低占比产品直接升级成普遍主推款，应先证明它在具体渠道或区域里的适配性。"
            )
        if target_profile:
            do_not_do.append(
                f"不要为了追 {target_profile['name']} 的总份额而平均撒网，优先级管理比“全面覆盖”更重要。"
            )
        if "价格字段" in missing_fields:
            do_not_do.append("不要在缺少稳定价格字段时直接制定大范围调价动作。")

        risk_controls = [
            "不要把一次静态分析当成长期趋势判断，缺时间字段时尤其如此。",
            "不要把白区直接理解为必投区域，高容量低份额只是进入候选，不是自动成立的商业结论。",
            "不要为了追求看起来完整的计划而同时推进过多对象，资源分散通常比机会不足更先导致失效。"
        ]
        if top3 >= 70:
            risk_controls.append("头部集中度较高，单一头部波动会放大整体结果波动，因此更需要把试点和守盘分开管理。")
        if price_insight is None:
            risk_controls.append("缺少稳定价格信息，任何涉及价格动作的建议都应先小范围验证。")

        return {
            'methodology': '本报告采用“诊断 -> 指导方针 -> 协调行动”的结构，强调先识别主要矛盾，再安排资源和动作顺序。',
            'executive_summary': executive_summary,
            'snapshot': snapshot,
            'primary_issue': primary_issue,
            'diagnosis_points': diagnosis_points,
            'model_sections': model_sections,
            'guiding_policies': guiding_policies,
            'priority_focus': priority_focus,
            'action_plan': action_plan,
            'do_not_do': do_not_do,
            'risk_controls': risk_controls,
            'target_profile': target_profile,
            'city_focus': city_focus,
            'channel_focus': channel_focus,
            'entity_focus': entity_focus,
            'product_focus': product_focus
        }

    def generate_word_report(self, analysis_data, chart_paths, report_title="数据分析报告"):
        """
        生成图文并茂的Word报告（轻量版，聚焦核心结论）
        """
        report_dir = self.output_root / 'reports'
        report_dir.mkdir(parents=True, exist_ok=True)

        doc = Document()
        doc.add_heading(report_title, 0)

        # 设置默认字体为微软雅黑
        style = doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        filtered_df = analysis_data['filtered_data']
        analysis_results = analysis_data['analysis_results']

        core_metric = analysis_results.get('核心指标列') or '核心指标'
        dim_label = analysis_results.get('核心维度') or '核心维度'
        display_dim_label = self.humanize_dimension_label(dim_label)
        blueprint = analysis_data.get('blueprint') or self.build_strategic_report_blueprint(filtered_df, analysis_results)

        doc.add_paragraph(blueprint['methodology'])

        doc.add_heading('执行摘要', level=1)
        for paragraph in blueprint['executive_summary']:
            doc.add_paragraph(paragraph)

        doc.add_heading('一、数据现状', level=1)
        for paragraph in blueprint['snapshot']:
            doc.add_paragraph(paragraph)

        doc.add_heading('二、核心诊断', level=1)
        doc.add_paragraph(blueprint['primary_issue'])
        for item in blueprint['diagnosis_points']:
            doc.add_paragraph(item, style='List Bullet')

        doc.add_heading('三、深度分析（思维模型）', level=1)
        for model in blueprint['model_sections']:
            doc.add_heading(model['title'], level=2)
            doc.add_paragraph(f"判断：{model['judgment']}")
            doc.add_paragraph(f"证据：{model['evidence']}")
            doc.add_paragraph(f"含义：{model['implication']}")
            if model.get('action'):
                doc.add_paragraph(f"建议动作：{model['action']}")

        doc.add_heading('四、指导方针', level=1)
        for policy in blueprint['guiding_policies']:
            doc.add_paragraph(policy, style='List Bullet')

        doc.add_heading('五、重点机会与优先级', level=1)
        doc.add_paragraph("评分逻辑：影响分衡量潜在增量，可行性分衡量落地基础，投入效率分衡量资源消耗可控性；综合优先级用于排序执行顺序。")
        has_priority = False
        if blueprint['priority_focus']['city']:
            has_priority = True
            doc.add_paragraph("城市优先清单：", style='List Bullet')
            for item in blueprint['priority_focus']['city']:
                doc.add_paragraph(item, style='List Number')
        if blueprint['priority_focus']['channel']:
            has_priority = True
            doc.add_paragraph("渠道优先清单：", style='List Bullet')
            for item in blueprint['priority_focus']['channel']:
                doc.add_paragraph(item, style='List Number')
        if blueprint['priority_focus']['entity']:
            has_priority = True
            doc.add_paragraph("重点实体优先清单：", style='List Bullet')
            for item in blueprint['priority_focus']['entity']:
                doc.add_paragraph(item, style='List Number')
        if not has_priority:
            doc.add_paragraph("当前数据不足以形成稳定的优先级清单，建议先补充目标对象份额与规模字段，再进入资源排序。", style='List Bullet')

        doc.add_heading('六、协调行动（90天）', level=1)
        for action in blueprint['action_plan']:
            doc.add_paragraph(action, style='List Number')

        doc.add_heading('七、不建议现在做的事', level=1)
        for item in blueprint['do_not_do']:
            doc.add_paragraph(item, style='List Bullet')

        doc.add_heading('八、现实约束与风险控制', level=1)
        for risk in blueprint['risk_controls']:
            doc.add_paragraph(risk, style='List Bullet')

        numeric_stats = analysis_results.get('数值列统计')
        concentration = analysis_results.get('集中度分析')
        if numeric_stats is not None and not numeric_stats.empty:
            doc.add_heading('附：关键数据口径（简版）', level=1)
            doc.add_paragraph(f"当前样本记录数 {len(filtered_df)}，字段数 {len(filtered_df.columns)}，核心观察维度为“{display_dim_label}”。")
            if concentration:
                doc.add_paragraph(
                    f"集中度参考：Top1 {concentration['Top1占比']:.1f}% / Top3 {concentration['Top3占比']:.1f}% / Top5 {concentration['Top5占比']:.1f}%。"
                )
            if core_metric in numeric_stats['字段'].values:
                row = numeric_stats[numeric_stats['字段'] == core_metric].iloc[0]
                doc.add_paragraph(
                    f"{core_metric}取值区间约为 {row['最小值']:,.2f} 到 {row['最大值']:,.2f}，说明不同对象体量差异较大，需要分层运营。"
                )

        # 嵌入图表（业务优先）
        if chart_paths:
            doc.add_heading('图表', level=1)
            def add_fig(key, caption):
                path_info = chart_paths.get(key)
                if isinstance(path_info, dict):
                    png = path_info.get('png')
                elif isinstance(path_info, (list, tuple)):
                    png = path_info[0]
                else:
                    png = None
                if png:
                    doc.add_paragraph(caption)
                    doc.add_picture(str(png), width=Inches(6.5))
            add_fig('core_share', f"{display_dim_label} 份额（按{core_metric or '指标'}）")
            add_fig('city_share', "城市分布（按量）")
            add_fig('category_share', "目录份额")
            add_fig('major_share', "产品大类分布")
            add_fig('coverage', f"{display_dim_label} 覆盖与单实体均量")
            add_fig('city_opportunities', "机会城市（如有）")
            add_fig('product_structure', "产品结构（如有）")
            add_fig('time_trend', "时间趋势（如有）")
            add_fig('correlation_heatmap', "相关性（如有）")

        invalid_chars = '<>:"/\\|?*'
        report_name = "".join('_' if ch in invalid_chars else ch for ch in report_title.replace(' ', '_')).strip(' .')
        if not report_name:
            report_name = "分析报告"
        report_path = report_dir / f"{report_name}.docx"
        doc.save(report_path)
        print(f"Word报告已生成: {report_path}")
        return report_path


    def run_full_pipeline(self, data_path, sheet_name=None, company_name=None):
        """
        运行完整的数据分析流水线

        Args:
            data_path: 数据文件路径
            sheet_name: Excel工作表名称
            company_name: 目标公司名称

        Returns:
            dict: 所有输出文件路径
        """
        print("="*60)
        print("开始执行完整的数据分析流水线")
        print("="*60)

        try:
            # 初始化输出目录与生成器
            self.output_root = self.resolve_output_root(data_path)
            self.setup_output_directories()
            self.init_generators()

            # 1. 数据分析
            analysis_data = self.analyze_data(data_path, sheet_name, company_name)

            # 2. 生成图表（按需）
            chart_paths = {}
            can_chart = self.should_generate_charts(analysis_data['analysis_results'])
            if can_chart:
                chart_paths = self.generate_charts(analysis_data['analysis_results'], company_name)

            # 3. 导出数据
            self.export_data(analysis_data['analysis_results'], analysis_data['insights'])

            # 4. 生成信息图（如启用图表则包含图片，否则空白图表区）
            html_path = self.generate_infographic(
                analysis_data['analysis_results'],
                chart_paths,
                analysis_data['insights'],
                company_name
            )

            # 4b. 生成图表汇总 HTML（便于快速预览/分发）
            gallery_path = self.create_chart_gallery(chart_paths)

            # 5. 生成截图（默认关闭）
            screenshot_path = None
            if self.config.get('enable_screenshot'):
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                screenshot_path = self.generate_screenshot(html_path, f"analysis_{timestamp}")

            # 6. 生成Word报告（聚焦核心结论）
            report_title = Path(data_path).stem + "_分析报告"
            word_report_path = self.generate_word_report(analysis_data, chart_paths, report_title)

            # 汇总所有输出
            outputs = {
                'charts': chart_paths,
                'html': html_path,
                'charts_gallery': gallery_path,
                'screenshot': screenshot_path,
                'word_report': str(word_report_path),
                'csv_dir': str(self.output_root / 'csv'),
                'excel': str(self.output_root / 'excel/数据分析汇总.xlsx'),
                'insights': str(self.output_root / 'reports/洞察分析.json'),
                'output_root': str(self.output_root)
            }

            print("\n" + "="*60)
            print("数据分析流水线执行完成！")
            print("="*60)
            print(f"信息图HTML: {html_path}")
            if screenshot_path:
                print(f"截图文件: {screenshot_path}")
            print(f"Word报告: {word_report_path}")
            print(f"图表文件: {len(chart_paths)}个")
            print(f"CSV文件: {self.output_root / 'csv'}")
            print(f"Excel汇总: {self.output_root / 'excel/数据分析汇总.xlsx'}")
            print(f"输出根目录: {self.output_root}")
            print("="*60)

            return outputs

        except Exception as e:
            print(f"流水线执行失败: {e}")
            import traceback
            traceback.print_exc()
            return None

def parse_arguments():
    """解析命令行参数"""
    parser = argparse.ArgumentParser(
        description="业务数据市场洞察分析包生成器"
    )
    parser.add_argument('data', help='数据文件路径（CSV/Excel）')
    parser.add_argument('legacy_sheet', nargs='?', default=None, help=argparse.SUPPRESS)
    parser.add_argument('legacy_company', nargs='?', default=None, help=argparse.SUPPRESS)
    parser.add_argument('--sheet', dest='sheet', help='Excel工作表名称')
    parser.add_argument('--company', dest='company', help='可选：需要高亮的品牌/分组名称')
    parser.add_argument('--time-column', dest='time_column', help='可选：时间列名称')
    parser.add_argument('--value-column', dest='value_column', help='可选：核心指标列名称')
    parser.add_argument('--output-dir', dest='output_dir', help='可选：输出目录（默认与数据同级的outputs/）')
    parser.add_argument('--enable-charts', dest='enable_charts', action='store_true', help='已废弃：请使用 --charts-mode on/off/auto（默认auto）')
    parser.add_argument('--enable-screenshot', dest='enable_screenshot', action='store_true', help='如需生成截图可开启，默认关闭')
    parser.add_argument('--charts-mode', dest='charts_mode', choices=['auto', 'on', 'off'], default='auto',
                        help='图表生成策略：auto(默认，只有数据有价值时生成)/on(强制生成)/off(关闭)')
    parser.add_argument('--core-dimension', dest='core_dimension', help='可选：核心实体维度列，如医院/客户/渠道/门店/品牌等')
    parser.add_argument('--target-brand', dest='target_brand', help='可选：目标品牌/申报企业，用于白区/机会分析')
    return parser.parse_args()


def main():
    """主函数"""
    args = parse_arguments()

    data_path = args.data
    sheet_name = args.sheet or args.legacy_sheet
    company_name = args.company if args.company is not None else args.legacy_company
    if not company_name and args.target_brand:
        company_name = args.target_brand

    if not os.path.exists(data_path):
        print(f"错误: 文件 {data_path} 不存在")
        sys.exit(1)

    config = {}
    if args.time_column:
        config['time_column'] = args.time_column
    if args.value_column:
        config['value_column'] = args.value_column
    if args.output_dir:
        config['output_dir'] = args.output_dir
    if args.enable_charts:
        config['charts_mode'] = 'on'
    if args.charts_mode:
        config['charts_mode'] = args.charts_mode
    if args.enable_screenshot:
        config['enable_screenshot'] = True
    if args.core_dimension:
        config['core_dimension'] = args.core_dimension
    if args.target_brand:
        config['target_brand'] = args.target_brand

    pipeline = DataAnalysisPipeline(config=config)
    results = pipeline.run_full_pipeline(data_path, sheet_name, company_name)

    if results:
        print("\n分析结果已保存到以下位置:")
        for key, value in results.items():
            if isinstance(value, dict):
                print(f"  {key}: {len(value)}个文件")
            else:
                print(f"  {key}: {value}")
    else:
        print("分析失败，请检查输入文件和参数")
        sys.exit(1)

if __name__ == "__main__":
    main()
