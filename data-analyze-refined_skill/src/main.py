#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
数据分析信息图主程序
整合所有模块，实现完整的数据分析信息图生成流程
"""

import os
import sys
import json
import argparse
import pandas as pd
from datetime import datetime
from pathlib import Path
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows

# 导入自定义模块
from data_analyzer import DataAnalyzer
from chart_generator import ChartGenerator
from infographic_generator import InfographicGenerator
from screenshot_generator import ScreenshotGenerator
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DataAnalysisPipeline:
    """数据分析流水线"""

    def __init__(self, config=None):
        """
        初始化流水线

        Args:
            config: 配置参数
        """
        self.config = config or {}
        self.analyzer = DataAnalyzer()
        self.chart_generator = None
        self.infographic_generator = None
        self.screenshot_generator = None
        self.output_root = None
        # 图表策略：auto/on/off
        self.charts_mode = self.config.get('charts_mode', 'auto')
        # 用户可指定的核心实体维度（如 医院/客户/渠道/门店/SKU 等）
        self.core_dimension = self.config.get('core_dimension')

    def resolve_output_root(self, data_path):
        """
        解析输出根目录：
        - 如果传入 output_dir，则使用传入值
        - 否则默认放在数据文件同级目录下的 outputs/
        """
        if self.config.get('output_dir'):
            return Path(self.config['output_dir']).expanduser().resolve()
        data_parent = Path(data_path).resolve().parent
        return data_parent / 'outputs'

    def setup_output_directories(self):
        """设置输出目录"""
        if self.output_root is None:
            raise ValueError("output_root 未初始化")
        directories = [
            self.output_root / 'csv',
            self.output_root / 'excel',
            self.output_root / 'figures' / 'png',
            self.output_root / 'figures' / 'svg',
            self.output_root / 'reports',
            self.output_root / 'html'
        ]

        for directory in directories:
            os.makedirs(directory, exist_ok=True)

    def init_generators(self):
        """基于输出目录初始化依赖组件"""
        if self.output_root is None:
            raise ValueError("output_root 未初始化")
        self.chart_generator = ChartGenerator(output_dir=str(self.output_root / 'figures'))
        self.infographic_generator = InfographicGenerator(
            template_dir='templates',
            output_dir=str(self.output_root / 'html')
        )
        # 截图默认关闭，若需要可后续按需初始化
        if self.config.get('enable_screenshot'):
            self.screenshot_generator = ScreenshotGenerator(
                output_dir=str(self.output_root / 'screenshots')
            )

    def should_generate_charts(self, analysis_results):
        """
        判断是否需要生成图表：
        - charts_mode == 'off' 时直接跳过
        - charts_mode == 'on' 时强制生成
        - charts_mode == 'auto' 时仅在有可视价值的板块才生成
        """
        if self.charts_mode == 'off':
            return False
        if self.charts_mode == 'on':
            return True

        # auto: 判定是否有可视价值
        numeric = analysis_results.get('数值列统计')
        cat = analysis_results.get('分类分布')
        corr = analysis_results.get('相关性矩阵')
        trend = analysis_results.get('时间趋势')
        brand = analysis_results.get('品牌份额')
        city = analysis_results.get('机会城市')
        coverage = analysis_results.get('覆盖分析')
        structure = analysis_results.get('产品结构')

        if numeric is not None and len(numeric) > 1:
            return True
        if cat is not None and not cat.empty and cat['数量'].max() > 1:
            return True
        if corr is not None:
            return True
        if trend is not None and not trend.empty and len(trend) > 2:
            return True
        if brand is not None and not brand.empty:
            return True
        if city is not None and not city.empty:
            return True
        if coverage is not None and not coverage.empty:
            return True
        if structure is not None and not structure.empty:
            return True
        return False

    def detect_data_files(self, data_dir='data'):
        """
        自动检测数据文件

        Args:
            data_dir: 数据目录

        Returns:
            list: 数据文件列表
        """
        data_files = []
        supported_extensions = ['.csv', '.xlsx', '.xls']

        if os.path.exists(data_dir):
            for file in os.listdir(data_dir):
                if any(file.endswith(ext) for ext in supported_extensions):
                    file_path = os.path.join(data_dir, file)
                    data_files.append(file_path)

        return data_files

    def build_market_narrative(
        self,
        dim_label,
        metric_label,
        concentration,
        head_names,
        tail_names,
        has_time,
        has_price,
        has_structure,
    ):
        """
        生成面向市场团队的叙事化摘要，避免学术统计口吻。
        """
        head = "、".join([str(x) for x in head_names[:3]]) if head_names else ""
        tail = "、".join([str(x) for x in tail_names[:3]]) if tail_names else ""

        if concentration:
            top3 = concentration.get("Top3占比", 0.0)
            top1 = concentration.get("Top1占比", 0.0)
            current = (
                f"当前{metric_label}主要由少数头部{dim_label}驱动，Top1约{top1:.1f}%，Top3约{top3:.1f}%。"
            )
        else:
            top3 = 0.0
            current = f"当前{metric_label}已出现明显头部与长尾分层，资源投入需要做优先级。"

        if head:
            current += f" 头部代表包括{head}。"

        opportunity = []
        if top3 >= 60:
            opportunity.append(f"头部{dim_label}是短期放量的直接抓手")
        if tail:
            opportunity.append(f"长尾中的{tail}适合作为低成本试点")
        if not opportunity:
            opportunity.append("可从覆盖不足和份额偏低的对象中筛选突破口")
        opportunity_text = "；".join(opportunity) + "。"

        strategy = (
            f"建议采用“头部深耕 + 白区试点”的双线策略：对头部{dim_label}做份额提升，对低份额高潜对象做定点突破。"
        )

        why = []
        why.append("这样做可以在不显著增加团队负担的情况下，把资源集中到最可能产生结果的对象上")
        if top3 >= 60:
            why.append("同时避免资源被低转化长尾平均摊薄")
        if not has_time:
            why.append("在缺少时间序列时，结构与覆盖策略比趋势判断更稳妥")
        why_text = "；".join(why) + "。"

        benefit = (
            f"预期收益是：更快形成可见增量、提升重点{dim_label}转化效率，并沉淀可复制打法扩展到相似对象。"
        )

        risks = []
        if not has_time:
            risks.append("缺少时间维度，短期效果评估可能滞后")
        if not has_price:
            risks.append("缺少价格维度，价格带策略判断存在盲区")
        if not has_structure:
            risks.append("缺少稳定结构字段，替代路径判断可能偏粗")
        if top3 >= 70:
            risks.append("头部集中度较高，单一对象波动会放大整体不确定性")
        if not risks:
            risks.append("当前主要风险来自执行节奏不一致")
        risk_text = "；".join(risks) + "。"

        mitigation = (
            "建议建立周节奏复盘：按对象跟踪进展、补齐关键字段（时间/价格/结构）、设置止损阈值并及时调整资源。"
        )

        return {
            "现状判断": current,
            "机会判断": opportunity_text,
            "策略建议": strategy,
            "为何现在做": why_text,
            "预期收益": benefit,
            "主要风险": risk_text,
            "风险对策": mitigation,
        }

    def analyze_data(self, data_path, sheet_name=None, company_name=None):
        """
        执行数据分析

        Args:
            data_path: 数据文件路径
            sheet_name: Excel工作表名称
            company_name: 目标公司名称

        Returns:
            dict: 分析结果
        """
        print(f"开始分析数据: {data_path}")

        # 1. 加载数据
        df = self.analyzer.load_data(data_path, sheet_name)
        df = self.analyzer.preprocess_data(df)
        print(f"数据加载完成，共{len(df)}条记录，{len(df.columns)}个字段")

        # 2. 数据体检
        health_report = self.analyzer.data_health_check(df)
        print("数据体检完成")

        # 3. 过滤留置针数据（如果适用）
        if '品种名称' in df.columns:
            df_filtered = self.analyzer.filter_iv_catheter(df)
            print(f"过滤后留置针数据: {len(df_filtered)}条记录")
        else:
            df_filtered = df

        # 4. 执行综合分析
        analysis_results = self.analyzer.create_comprehensive_analysis(
            df_filtered, company_name, self.config
        )
        print("综合分析完成")

        # 5. 生成洞察
        insights = self.analyzer.generate_insights(analysis_results, company_name)

        return {
            'raw_data': df,
            'filtered_data': df_filtered,
            'health_report': health_report,
            'analysis_results': analysis_results,
            'insights': insights
        }

    def generate_charts(self, analysis_results, company_name=None):
        """
        生成可视化图表

        Args:
            analysis_results: 分析结果

        Returns:
            dict: 图表路径
        """
        print("开始生成图表...")
        chart_paths = self.chart_generator.generate_all_charts(
            analysis_results, company_name=company_name or ''
        )
        print(f"图表生成完成，共{len(chart_paths)}个图表")
        return chart_paths

    def export_data(self, analysis_results, insights):
        """
        导出分析数据

        Args:
            analysis_results: 分析结果
            insights: 洞察分析
        """
        print("开始导出数据...")

        csv_dir = self.output_root / 'csv'
        # 导出CSV文件
        csv_exports = {
            '字段概览': analysis_results.get('字段概览'),
            '数值列统计': analysis_results.get('数值列统计'),
            '分类分布': analysis_results.get('分类分布'),
            '相关性矩阵': analysis_results.get('相关性矩阵'),
            '时间趋势': analysis_results.get('时间趋势'),
            '核心维度分布': analysis_results.get('核心维度分布'),
            '城市分布': analysis_results.get('城市分布'),
            '城市品牌分布': analysis_results.get('城市品牌分布'),
            '城市Top3': analysis_results.get('城市Top3'),
            '城市白区': analysis_results.get('城市白区'),
            '机会优先级_城市': analysis_results.get('机会优先级_城市'),
            '机会城市': analysis_results.get('机会城市'),
            '机会医院': analysis_results.get('机会医院'),
            '目录分布': analysis_results.get('目录分布'),
            '大类分布': analysis_results.get('大类分布'),
            '产品结构': analysis_results.get('产品结构'),
            '覆盖分析': analysis_results.get('覆盖分析'),
            '医院TOP': analysis_results.get('医院TOP'),
            '产品分布': analysis_results.get('产品分布'),
            '品牌产品分布': analysis_results.get('品牌产品分布'),
            '城市产品分布': analysis_results.get('城市产品分布'),
            '品牌产品Top': analysis_results.get('品牌产品Top'),
            '医院白院': analysis_results.get('医院白院'),
            '机会优先级_医院': analysis_results.get('机会优先级_医院')
        }

        for name, data in csv_exports.items():
            if data is not None and not data.empty:
                filepath = csv_dir / f"{name}.csv"
                data.to_csv(filepath, index=False, encoding='utf-8-sig')
                print(f"导出CSV: {filepath}")

        # 导出Excel汇总表（带目录，移除字段概览/数值统计）
        excel_path = self.output_root / "excel/数据分析汇总.xlsx"
        wb = Workbook()
        toc = wb.active
        toc.title = "目录"
        toc.append(["名称", "跳转"])

        def add_toc_row(name):
            row = toc.max_row + 1
            toc.cell(row=row, column=1, value=name)
            cell = toc.cell(row=row, column=2, value="点击跳转")
            cell.hyperlink = f"#{name}!A1"
            cell.style = "Hyperlink"

        exclude_excel = {"字段概览", "数值列统计"}
        for name, data in csv_exports.items():
            if name in exclude_excel or data is None or data.empty:
                continue
            ws = wb.create_sheet(name)
            for r in dataframe_to_rows(data, index=False, header=True):
                ws.append(r)
            add_toc_row(name)
            back = ws.cell(row=ws.max_row + 2, column=1)
            back.value = "返回目录"
            back.hyperlink = "#目录!A1"
            back.style = "Hyperlink"

        if "Sheet" in wb.sheetnames:
            wb.remove(wb["Sheet"])
        wb.save(excel_path)
        print(f"导出Excel: {excel_path}")

        # 导出洞察分析
        insights_path = self.output_root / "reports/洞察分析.json"
        with open(insights_path, 'w', encoding='utf-8') as f:
            json.dump(insights, f, ensure_ascii=False, indent=2)
        print(f"导出洞察: {insights_path}")

    def generate_infographic(self, analysis_results, chart_paths, insights, company_name=None):
        """
        生成信息图HTML

        Args:
            analysis_results: 分析结果
            chart_paths: 图表路径
            insights: 洞察分析
            company_name: 公司名称

        Returns:
            str: HTML文件路径
        """
        print("生成信息图HTML...")
        html_content = self.infographic_generator.generate_infographic_html(
            analysis_results, chart_paths, insights, company_name
        )

        html_path = self.infographic_generator.save_infographic(html_content)
        print(f"信息图HTML已生成: {html_path}")
        return html_path

    def generate_screenshot(self, html_path, output_name=None):
        """按需生成截图（默认关闭）"""
        if not self.config.get('enable_screenshot') or not self.screenshot_generator:
            return None
        print("生成截图...")
        screenshot_path = self.screenshot_generator.generate_screenshot_sync(
            html_path, output_name
        )
        print(f"截图已生成: {screenshot_path}")
        return screenshot_path

    def create_chart_gallery(self, chart_paths):
        """生成图表汇总 HTML，方便一次性查看/下载"""
        if not chart_paths:
            return None
        html_dir = self.output_root / 'html'
        html_dir.mkdir(parents=True, exist_ok=True)
        mapping = {
            'core_share': '核心维度份额',
            'city_share': '城市分布',
            'category_share': '目录份额',
            'major_share': '产品大类分布',
            'coverage': '覆盖与单院均量',
            'city_opportunities': '机会城市',
            'product_structure': '产品结构',
            'time_trend': '时间趋势',
            'correlation_heatmap': '相关性',
            'categorical_topn': '分类TopN',
            'numeric_overview': '数值概览',
        }
        items = []
        from pathlib import Path
        for key, val in chart_paths.items():
            png = None
            if isinstance(val, dict):
                png = val.get('png')
            elif isinstance(val, (list, tuple)) and len(val) > 0:
                png = val[0]
            if png and Path(png).exists():
                items.append((mapping.get(key, key), png))
        if not items:
            return None

        html_parts = [
            "<!DOCTYPE html>",
            "<html lang='zh-CN'>",
            "<head>",
            "<meta charset='UTF-8' />",
            "<title>图表汇总</title>",
            "<style>",
            "body { font-family: Arial, sans-serif; margin: 24px; background: #f7f7f7; }",
            "h1 { margin-bottom: 12px; }",
            ".card { background: #fff; padding: 16px; margin-bottom: 16px; box-shadow: 0 2px 6px rgba(0,0,0,0.08); border-radius: 8px; }",
            ".card h2 { margin: 0 0 8px 0; font-size: 18px; }",
            ".card img { max-width: 100%; height: auto; border: 1px solid #e5e5e5; border-radius: 4px; }",
            "</style>",
            "</head>",
            "<body>",
            "<h1>图表汇总</h1>",
        ]
        for title, png in items:
            html_parts.append("<div class='card'>")
            html_parts.append(f"<h2>{title}</h2>")
            html_parts.append(f"<img src='file://{png}' alt='{title}' />")
            html_parts.append("</div>")
        html_parts.append("</body></html>")
        out_path = html_dir / "charts_gallery.html"
        out_path.write_text("\n".join(html_parts), encoding="utf-8")
        print(f"图表汇总已生成: {out_path}")
        return str(out_path)

    def generate_word_report(self, analysis_data, chart_paths, report_title="数据分析报告"):
        """
        生成图文并茂的Word报告（轻量版，聚焦核心结论）
        """
        report_dir = self.output_root / 'reports'
        report_dir.mkdir(parents=True, exist_ok=True)

        doc = Document()
        doc.add_heading(report_title, 0)

        # 设置默认字体为微软雅黑
        style = doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        filtered_df = analysis_data['filtered_data']
        analysis_results = analysis_data['analysis_results']

        # 基础信息与关键维度
        core_metric = analysis_results.get('核心指标列')
        numeric_stats = analysis_results.get('数值列统计')
        core_dim = analysis_results.get('核心维度') or '核心维度'
        dim_label = core_dim
        concentration = analysis_results.get('集中度分析')
        head_names = analysis_results.get('头部名单') or []
        tail_names = analysis_results.get('尾部名单') or []
        city_priority = analysis_results.get('机会优先级_城市')
        hosp_priority = analysis_results.get('机会优先级_医院')

        def format_entities(names, max_n=5):
            names = [
                str(n).strip()
                for n in names
                if pd.notna(n) and str(n).strip() and str(n).strip().lower() not in {'nan', 'none', 'null', 'na', 'n/a'}
            ]
            if not names:
                return ""
            subset = names[:max_n]
            if len(names) > max_n:
                return "、".join(subset) + " 等"
            return "、".join(subset)

        head_names_fmt = format_entities(head_names)
        tail_names_fmt = format_entities(tail_names)

        # 价格-量关系分析（如有价格字段）
        price_info = analysis_results.get('价格分析') or {}
        price_col = price_info.get('价格列')
        price_insight = None
        price_quadrant_text = None
        if price_info:
            corr = price_info.get('相关系数')
            relation = price_info.get('相关描述')
            if corr is not None and relation:
                price_insight = f"{price_col} 与 {core_metric} 的相关系数约 {corr:.2f}，{relation}"
            lphv_names = price_info.get('低价高量') or []
            hplv_names = price_info.get('高价低量') or []
            segments = []
            if lphv_names:
                segments.append(f"低价高量代表{dim_label}如：{format_entities(lphv_names)}")
            if hplv_names:
                segments.append(f"高价低量{dim_label}如：{format_entities(hplv_names)}")
            if segments:
                price_quadrant_text = "；".join(segments) + "。"

        # 结构机会（如有产品/类别结构分析）
        structure_insight = None
        structure_summary = analysis_results.get('结构概览')
        if structure_summary:
            main_type = structure_summary.get('主类型')
            main_share = structure_summary.get('主类型占比')
            others = structure_summary.get('其他占比')
            if main_type is not None and main_share is not None:
                if others:
                    structure_insight = f"当前业务量主要来自 {main_type}（约 {main_share:.1f}%），其余类型为：{others}"
                else:
                    structure_insight = f"当前业务量几乎全部来自 {main_type}（约 {main_share:.1f}%），结构相对单一"

        # 时间趋势（若存在）
        time_trend = analysis_results.get('时间趋势')
        trend_summary_obj = analysis_results.get('时间趋势摘要')
        time_trend_summary = None
        if trend_summary_obj:
            start_v = trend_summary_obj.get('起始')
            end_v = trend_summary_obj.get('当前')
            dir_text = trend_summary_obj.get('方向')
            if start_v is not None and end_v is not None and dir_text:
                time_trend_summary = f"整体时间趋势呈{dir_text}态势（起始值 {start_v:,.1f} → 当前值 {end_v:,.1f}）"

        narrative = self.build_market_narrative(
            dim_label=dim_label,
            metric_label=core_metric or '核心指标',
            concentration=concentration,
            head_names=head_names,
            tail_names=tail_names,
            has_time=bool(time_trend_summary),
            has_price=bool(price_insight),
            has_structure=bool(structure_insight),
        )

        # 开场使用市场叙事，不先堆统计口径
        doc.add_heading('执行摘要（市场视角）', level=1)
        doc.add_paragraph(f"现状判断：{narrative['现状判断']}")
        doc.add_paragraph(f"机会判断：{narrative['机会判断']}")
        doc.add_paragraph(f"策略建议：{narrative['策略建议']}")
        doc.add_paragraph(f"为何现在做：{narrative['为何现在做']}")
        doc.add_paragraph(f"预期收益：{narrative['预期收益']}")

        doc.add_heading('现状与机会解读', level=1)
        if concentration:
            doc.add_paragraph(
                f"从结构上看，前3个{dim_label}已覆盖约 {concentration['Top3占比']:.1f}% 的{core_metric or '业务量'}，说明当前增长更依赖头部对象。"
            )
        if head_names_fmt:
            doc.add_paragraph(
                f"头部对象（{head_names_fmt}）是短期增量最确定的来源，优先投入会比平均分配更高效。"
            )
        if tail_names_fmt:
            doc.add_paragraph(
                f"尾部对象（{tail_names_fmt}）适合作为低成本试点：一旦跑通，可复制到同类对象，形成第二增长曲线。"
            )
        if structure_insight:
            doc.add_paragraph(f"结构层面：{structure_insight}。这意味着应先围绕高占比结构做深做透，再扩展到次要结构。")
        if time_trend_summary:
            doc.add_paragraph(f"时间维度：{time_trend_summary}。可据此安排阶段性活动节奏和复盘节点。")

        doc.add_heading('策略路径与理由', level=1)
        doc.add_paragraph(narrative['策略建议'])
        doc.add_paragraph(narrative['为何现在做'])
        if price_insight:
            text = price_insight
            if price_quadrant_text:
                text += "；" + price_quadrant_text
            doc.add_paragraph(f"价格与量关系补充：{text}")

        # 机会优先级（影响 × 可行性 × 投入）
        doc.add_heading('机会优先级（影响×可行性×投入）', level=1)
        doc.add_paragraph("评分口径：影响分衡量潜在增量，可行性分衡量落地基础，投入效率分衡量资源消耗可控性；综合优先级用于排序执行顺序。")
        has_priority = False

        if city_priority is not None and not city_priority.empty:
            has_priority = True
            doc.add_paragraph("城市机会 TOP（建议优先推进）：", style='List Bullet')
            show_cols = ['城市', '城市总量', '目标品牌份额(%)', '综合优先级分', '优先级', '优先级理由']
            city_show = city_priority[[c for c in show_cols if c in city_priority.columns]].head(5)
            for _, r in city_show.iterrows():
                name = r.get('城市', '未知城市')
                total = r.get('城市总量', 0)
                share = r.get('目标品牌份额(%)', 0)
                score = r.get('综合优先级分', 0)
                level = r.get('优先级', '')
                reason = r.get('优先级理由', '')
                doc.add_paragraph(
                    f"{name}：优先级 {level}（{score:.1f}分），当前份额约 {share:.1f}%，容量约 {total:,.0f}。理由：{reason}。",
                    style='List Number'
                )

        if hosp_priority is not None and not hosp_priority.empty:
            has_priority = True
            entity_col = '医院名称'
            for cand in ['医院名称', '医疗机构名称']:
                if cand in hosp_priority.columns:
                    entity_col = cand
                    break
            doc.add_paragraph("机构机会 TOP（建议作为试点突破）：", style='List Bullet')
            show_cols = [entity_col, '医院总量', '目标品牌份额(%)', '综合优先级分', '优先级', '优先级理由']
            hosp_show = hosp_priority[[c for c in show_cols if c in hosp_priority.columns]].head(5)
            for _, r in hosp_show.iterrows():
                name = r.get(entity_col, '未知机构')
                total = r.get('医院总量', 0)
                share = r.get('目标品牌份额(%)', 0)
                score = r.get('综合优先级分', 0)
                level = r.get('优先级', '')
                reason = r.get('优先级理由', '')
                doc.add_paragraph(
                    f"{name}：优先级 {level}（{score:.1f}分），当前份额约 {share:.1f}%，容量约 {total:,.0f}。理由：{reason}。",
                    style='List Number'
                )

        if not has_priority:
            doc.add_paragraph("当前数据不足以生成稳定的优先级清单，建议补充目标对象份额与规模字段后再进行排序。", style='List Bullet')

        if numeric_stats is not None and not numeric_stats.empty:
            doc.add_heading('支撑数据（简版）', level=1)
            doc.add_paragraph(f"当前样本记录数 {len(filtered_df)}，字段数 {len(filtered_df.columns)}，核心观察维度为“{dim_label}”。")
            if concentration:
                doc.add_paragraph(
                    f"集中度参考：Top1 {concentration['Top1占比']:.1f}% / Top3 {concentration['Top3占比']:.1f}% / Top5 {concentration['Top5占比']:.1f}%。"
                )
            if core_metric and core_metric in numeric_stats['字段'].values:
                row = numeric_stats[numeric_stats['字段'] == core_metric].iloc[0]
                doc.add_paragraph(
                    f"{core_metric}取值区间约为 {row['最小值']:,.2f} 到 {row['最大值']:,.2f}，说明不同对象体量差异较大，需要分层运营。"
                )

        # 风险分析
        doc.add_heading('风险分析', level=1)
        doc.add_paragraph(f"主要风险：{narrative['主要风险']}")
        doc.add_paragraph(f"风险对策：{narrative['风险对策']}")

        # 行动清单
        doc.add_heading('下一步行动计划', level=1)
        actions = []
        if head_names_fmt:
            actions.append(
                f"生成头部{dim_label}清单（如 {head_names_fmt}），为每家指定负责人与季度目标，优先落实客情、准入和配额。"
            )
        if tail_names_fmt:
            actions.append(
                f"尾部试点：从尾部{dim_label}中筛选少量试点对象（如 {tail_names_fmt}），设计轻量试用/准入项目，其余暂不投入。"
            )
        if price_insight:
            actions.append(
                f"针对高价低量{dim_label}逐一评估是否需要调价/替代或增加价值服务；对低价高量{dim_label}梳理打法，形成可复制 SOP。"
            )
        if structure_insight:
            actions.append(
                "围绕占比较高的结构类型设计重点推广物料和项目，同时规划其他结构的切入节奏与试点场景。"
            )
        missing_fields = []
        for col in ['时间', '日期', '月份', price_col, '渠道', '渠道名称']:
            if col and col not in filtered_df.columns:
                missing_fields.append(col)
        if missing_fields:
            actions.append(
                "在后续数据采集里补充关键字段（如 {}），以便做趋势/渠道贡献度/价格敏感度等更完整的分析。".format(
                    "、".join(sorted(set([c for c in missing_fields if c])))
                )
            )
        if not actions:
            actions.append("基于当前分析结果，组织一次复盘会，确认头部名单、试点对象和下一个评估节点。")

        for a in actions:
            doc.add_paragraph(a, style='List Number')

        # 嵌入图表（业务优先）
        if chart_paths:
            doc.add_heading('图表', level=1)
            def add_fig(key, caption):
                path_info = chart_paths.get(key)
                if isinstance(path_info, dict):
                    png = path_info.get('png')
                elif isinstance(path_info, (list, tuple)):
                    png = path_info[0]
                else:
                    png = None
                if png:
                    doc.add_paragraph(caption)
                    doc.add_picture(str(png), width=Inches(6.5))
            add_fig('core_share', f"{dim_label} 份额（按{core_metric or '指标'}）")
            add_fig('city_share', "城市分布（按量）")
            add_fig('category_share', "目录份额")
            add_fig('major_share', "产品大类分布")
            add_fig('coverage', f"{dim_label} 覆盖与单院均量")
            add_fig('city_opportunities', "机会城市（如有）")
            add_fig('product_structure', "产品结构（如有）")
            add_fig('time_trend', "时间趋势（如有）")
            add_fig('correlation_heatmap', "相关性（如有）")

        report_name = report_title.replace(' ', '_')
        report_path = report_dir / f"{report_name}.docx"
        doc.save(report_path)
        print(f"Word报告已生成: {report_path}")
        return report_path


    def run_full_pipeline(self, data_path, sheet_name=None, company_name=None):
        """
        运行完整的数据分析流水线

        Args:
            data_path: 数据文件路径
            sheet_name: Excel工作表名称
            company_name: 目标公司名称

        Returns:
            dict: 所有输出文件路径
        """
        print("="*60)
        print("开始执行完整的数据分析流水线")
        print("="*60)

        try:
            # 初始化输出目录与生成器
            self.output_root = self.resolve_output_root(data_path)
            self.setup_output_directories()
            self.init_generators()

            # 1. 数据分析
            analysis_data = self.analyze_data(data_path, sheet_name, company_name)

            # 2. 生成图表（按需）
            chart_paths = {}
            can_chart = self.should_generate_charts(analysis_data['analysis_results'])
            if can_chart:
                chart_paths = self.generate_charts(analysis_data['analysis_results'], company_name)

            # 3. 导出数据
            self.export_data(analysis_data['analysis_results'], analysis_data['insights'])

            # 4. 生成信息图（如启用图表则包含图片，否则空白图表区）
            html_path = self.generate_infographic(
                analysis_data['analysis_results'],
                chart_paths,
                analysis_data['insights'],
                company_name
            )

            # 4b. 生成图表汇总 HTML（便于快速预览/分发）
            gallery_path = self.create_chart_gallery(chart_paths)

            # 5. 生成截图（默认关闭）
            screenshot_path = None
            if self.config.get('enable_screenshot'):
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                screenshot_path = self.generate_screenshot(html_path, f"analysis_{timestamp}")

            # 6. 生成Word报告（聚焦核心结论）
            report_title = Path(data_path).stem + "_分析报告"
            word_report_path = self.generate_word_report(analysis_data, chart_paths, report_title)

            # 汇总所有输出
            outputs = {
                'charts': chart_paths,
                'html': html_path,
                'charts_gallery': gallery_path,
                'screenshot': screenshot_path,
                'word_report': str(word_report_path),
                'csv_dir': str(self.output_root / 'csv'),
                'excel': str(self.output_root / 'excel/数据分析汇总.xlsx'),
                'insights': str(self.output_root / 'reports/洞察分析.json'),
                'output_root': str(self.output_root)
            }

            print("\n" + "="*60)
            print("数据分析流水线执行完成！")
            print("="*60)
            print(f"信息图HTML: {html_path}")
            if screenshot_path:
                print(f"截图文件: {screenshot_path}")
            print(f"Word报告: {word_report_path}")
            print(f"图表文件: {len(chart_paths)}个")
            print(f"CSV文件: {self.output_root / 'csv'}")
            print(f"Excel汇总: {self.output_root / 'excel/数据分析汇总.xlsx'}")
            print(f"输出根目录: {self.output_root}")
            print("="*60)

            return outputs

        except Exception as e:
            print(f"流水线执行失败: {e}")
            import traceback
            traceback.print_exc()
            return None

def parse_arguments():
    """解析命令行参数"""
    parser = argparse.ArgumentParser(
        description="通用数据分析信息图生成器"
    )
    parser.add_argument('data', help='数据文件路径（CSV/Excel）')
    parser.add_argument('legacy_sheet', nargs='?', default=None, help=argparse.SUPPRESS)
    parser.add_argument('legacy_company', nargs='?', default=None, help=argparse.SUPPRESS)
    parser.add_argument('--sheet', dest='sheet', help='Excel工作表名称')
    parser.add_argument('--company', dest='company', help='可选：需要高亮的品牌/分组名称')
    parser.add_argument('--time-column', dest='time_column', help='可选：时间列名称')
    parser.add_argument('--value-column', dest='value_column', help='可选：核心指标列名称')
    parser.add_argument('--output-dir', dest='output_dir', help='可选：输出目录（默认与数据同级的outputs/）')
    parser.add_argument('--enable-charts', dest='enable_charts', action='store_true', help='已废弃：请使用 --charts-mode on/off/auto（默认auto）')
    parser.add_argument('--enable-screenshot', dest='enable_screenshot', action='store_true', help='如需生成截图可开启，默认关闭')
    parser.add_argument('--charts-mode', dest='charts_mode', choices=['auto', 'on', 'off'], default='auto',
                        help='图表生成策略：auto(默认，只有数据有价值时生成)/on(强制生成)/off(关闭)')
    parser.add_argument('--core-dimension', dest='core_dimension', help='可选：核心实体维度列，如医院/客户/渠道/门店/品牌等')
    parser.add_argument('--target-brand', dest='target_brand', help='可选：目标品牌/申报企业，用于白区/机会分析')
    return parser.parse_args()


def main():
    """主函数"""
    args = parse_arguments()

    data_path = args.data
    sheet_name = args.sheet or args.legacy_sheet
    company_name = args.company if args.company is not None else args.legacy_company

    if not os.path.exists(data_path):
        print(f"错误: 文件 {data_path} 不存在")
        sys.exit(1)

    config = {}
    if args.time_column:
        config['time_column'] = args.time_column
    if args.value_column:
        config['value_column'] = args.value_column
    if args.output_dir:
        config['output_dir'] = args.output_dir
    if args.enable_charts:
        config['charts_mode'] = 'on'
    if args.charts_mode:
        config['charts_mode'] = args.charts_mode
    if args.enable_screenshot:
        config['enable_screenshot'] = True
    if args.core_dimension:
        config['core_dimension'] = args.core_dimension
    if args.target_brand:
        config['target_brand'] = args.target_brand

    pipeline = DataAnalysisPipeline(config=config)
    results = pipeline.run_full_pipeline(data_path, sheet_name, company_name)

    if results:
        print("\n分析结果已保存到以下位置:")
        for key, value in results.items():
            if isinstance(value, dict):
                print(f"  {key}: {len(value)}个文件")
            else:
                print(f"  {key}: {value}")
    else:
        print("分析失败，请检查输入文件和参数")
        sys.exit(1)

if __name__ == "__main__":
    main()
